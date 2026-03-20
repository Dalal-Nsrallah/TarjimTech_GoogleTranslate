"""
مُصدّر الملفات (Exporter)
===========================
يُصدّر النص النهائي والتقارير بعدة صيغ:
  - TXT: نص عادي
  - DOCX: ملف Word مع تنسيق عربي
  - JSON: التقرير الكامل بكل التفاصيل
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Optional


class Exporter:
    """مُصدّر النتائج بصيغ مختلفة."""

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _generate_filename(self, prefix: str, ext: str) -> Path:
        """توليد اسم ملف فريد بالتاريخ والوقت."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return self.output_dir / f"{prefix}_{timestamp}.{ext}"

    # ═══════════════════════════════════════════
    # تصدير TXT
    # ═══════════════════════════════════════════

    def to_txt(self, pipeline_result, filename: Optional[str] = None) -> str:
        """
        تصدير النص النهائي كملف TXT.

        Args:
            pipeline_result: نتيجة خط الأنابيب
            filename: اسم الملف (اختياري)

        Returns:
            مسار الملف المحفوظ
        """
        if filename:
            path = self.output_dir / filename
        else:
            path = self._generate_filename("translation", "txt")

        lines = []
        lines.append("=" * 60)
        lines.append("  TarjimTech — نتيجة الترجمة المحسّنة")
        lines.append(f"  التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("=" * 60)
        lines.append("")

        # النص المصدر
        lines.append("── النص المصدر (الإنجليزي) ──")
        lines.append(pipeline_result.source)
        lines.append("")

        # النص الأصلي
        lines.append("── الترجمة الأصلية ──")
        lines.append(pipeline_result.target_original)
        lines.append("")

        # النص النهائي
        lines.append("── النص النهائي المحسّن ──")
        lines.append(pipeline_result.final_text)
        lines.append("")

        # تقرير الجودة
        lines.append("─" * 40)
        lines.append(f"  درجة الجودة: {pipeline_result.quality_score:.1f}/100")
        lines.append(f"  التصنيف: {pipeline_result.quality_grade}")
        lines.append("")

        # تحسينات الفصاحة
        if pipeline_result.eloquence_fixes:
            lines.append("── تحسينات الفصاحة ──")
            for fix in pipeline_result.eloquence_fixes:
                lines.append(f"  • {fix['original']} → {fix['improved']}")
                lines.append(f"    السبب: {fix['explanation']}")
            lines.append("")
            lines.append(
                f"  فصاحة: {pipeline_result.eloquence_score_before:.0f} → "
                f"{pipeline_result.eloquence_score_after:.0f}"
            )
            lines.append("")

        # مشاكل الجودة
        if pipeline_result.quality_issues:
            lines.append("── ملاحظات الجودة ──")
            for issue in pipeline_result.quality_issues:
                status = "✓ مُصلح" if issue.get("applied") else "⚠ ملاحظة"
                lines.append(f"  [{status}] {issue['description']}")
                if issue.get("suggestion"):
                    lines.append(f"           اقتراح: {issue['suggestion']}")
            lines.append("")

        lines.append("=" * 60)
        lines.append("  تم التوليد بواسطة TarjimTech — نظام الوكلاء الستة")
        lines.append("=" * 60)

        content = "\n".join(lines)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        return str(path)

    # ═══════════════════════════════════════════
    # تصدير DOCX (Word)
    # ═══════════════════════════════════════════

    def to_docx(self, pipeline_result, filename: Optional[str] = None) -> str:
        """
        تصدير النص النهائي كملف Word مع تنسيق عربي.

        Args:
            pipeline_result: نتيجة خط الأنابيب
            filename: اسم الملف (اختياري)

        Returns:
            مسار الملف المحفوظ
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # إذا python-docx مو مثبتة، صدّر كـ TXT بامتداد .docx.txt
            alt_path = self._generate_filename("translation", "docx.txt")
            txt_path = self.to_txt(pipeline_result,
                                   filename=alt_path.name)
            return txt_path + "\n⚠ لتصدير Word: pip install python-docx"

        if filename:
            path = self.output_dir / filename
        else:
            path = self._generate_filename("translation", "docx")

        doc = Document()

        # ── إعدادات الخط الافتراضي ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(13)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ── العنوان ──
        title = doc.add_heading("", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = title.add_run("TarjimTech — نتيجة الترجمة المحسّنة")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_para.add_run(
            f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # ── النص المصدر ──
        h = doc.add_heading("", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = h.add_run("النص المصدر (الإنجليزي)")
        run.font.color.rgb = RGBColor(0x42, 0x85, 0xF4)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # إنجليزي = يسار
        run = p.add_run(pipeline_result.source)
        run.font.size = Pt(12)

        # ── الترجمة الأصلية ──
        h = doc.add_heading("", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = h.add_run("الترجمة الأصلية")
        run.font.color.rgb = RGBColor(0xEA, 0x43, 0x35)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(pipeline_result.target_original)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

        # ── النص النهائي المحسّن ──
        h = doc.add_heading("", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = h.add_run("النص النهائي المحسّن ✓")
        run.font.color.rgb = RGBColor(0x34, 0xA8, 0x53)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(pipeline_result.final_text)
        run.font.size = Pt(14)
        run.bold = True

        doc.add_paragraph("")

        # ── تقرير الجودة ──
        h = doc.add_heading("", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = h.add_run("تقرير الجودة")

        table = doc.add_table(rows=2, cols=2)
        table.style = "Light Grid Accent 1"

        table.cell(0, 0).text = "التصنيف"
        table.cell(0, 1).text = pipeline_result.quality_grade
        table.cell(1, 0).text = "الدرجة"
        score = pipeline_result.quality_final_score or pipeline_result.quality_score
        table.cell(1, 1).text = f"{score:.1f}/100"

        # محاذاة الجدول لليمين
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("")

        # ── تحسينات الفصاحة ──
        if pipeline_result.eloquence_fixes:
            h = doc.add_heading("", level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = h.add_run("تحسينات الفصاحة")
            run.font.color.rgb = RGBColor(0xFB, 0xBC, 0x04)

            for fix in pipeline_result.eloquence_fixes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                run = p.add_run(f"{fix['original']}")
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.strikethrough = True

                run = p.add_run(f"  →  ")

                run = p.add_run(f"{fix['improved']}")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run.bold = True

                exp = doc.add_paragraph()
                exp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = exp.add_run(f"    {fix['explanation']}")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(
                f"الفصاحة: {pipeline_result.eloquence_score_before:.0f} → "
                f"{pipeline_result.eloquence_score_after:.0f}"
            )
            run.bold = True

        # ── التذييل ──
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("تم التوليد بواسطة TarjimTech — نظام الوكلاء الستة")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

        doc.save(str(path))
        return str(path)

    # ═══════════════════════════════════════════
    # تصدير JSON (التقرير الكامل)
    # ═══════════════════════════════════════════

    def to_json(self, pipeline_result, filename: Optional[str] = None) -> str:
        """تصدير التقرير الكامل كـ JSON."""
        if filename:
            path = self.output_dir / filename
        else:
            path = self._generate_filename("report", "json")

        data = pipeline_result.to_dict()
        data["exported_at"] = datetime.now().isoformat()

        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        return str(path)

    # ═══════════════════════════════════════════
    # تصدير الكل دفعة واحدة
    # ═══════════════════════════════════════════

    def export_all(self, pipeline_result, base_name: Optional[str] = None) -> dict:
        """
        تصدير بكل الصيغ دفعة واحدة.

        Returns:
            {"txt": path, "docx": path, "json": path}
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        prefix = base_name or f"translation_{timestamp}"

        return {
            "txt": self.to_txt(pipeline_result, f"{prefix}.txt"),
            "docx": self.to_docx(pipeline_result, f"{prefix}.docx"),
            "json": self.to_json(pipeline_result, f"{prefix}.json"),
        }
